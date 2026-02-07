#!/usr/bin/env python3
"""Convert QA_CHECKLIST.md to QA_CHECKLIST.docx (Word), saved to Desktop."""

import re
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DESKTOP = Path.home() / "Desktop"
MD_PATH = Path(__file__).resolve().parent.parent / "QA_CHECKLIST.md"
OUT_PATH = DESKTOP / "QA_CHECKLIST.docx"


def strip_markdown_to_plain(s):
    """Convert markdown to plain text: **bold** -> bold, `code` -> code, [text](url) -> text."""
    if not s:
        return s
    # [link text](url) -> link text
    s = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", s)
    # `code` -> code
    s = re.sub(r"`([^`]+)`", r"\1", s)
    # **bold** -> bold (remove markers only; Word run bold applied separately)
    return s


def set_paragraph_formatted_text(paragraph, text):
    """Set paragraph content, converting **bold** to Word bold runs. Clears existing content."""
    # Split by **...** so odd-index segments are bold; then strip ` and []( ) from each part
    parts = re.split(r"\*\*(.+?)\*\*", text)
    paragraph.clear()
    for i, part in enumerate(parts):
        if not part:
            continue
        plain = strip_markdown_to_plain(part)
        run = paragraph.add_run(plain)
        if i % 2 == 1:
            run.bold = True


def add_table_from_md_lines(doc, lines):
    """Parse markdown table lines and add a Word table. Skip separator row (----)."""
    rows_data = []
    for line in lines:
        line = line.strip()
        if not line or not line.startswith("|"):
            continue
        parts = [p.strip() for p in line.split("|")]
        parts = [p for p in parts if p]  # remove empty from leading/trailing |
        if not parts:
            continue
        if all(c in "-" for c in parts[0]):  # separator row
            continue
        rows_data.append(parts)
    if not rows_data:
        return
    num_cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.style = "Table Grid"
    for i, row_data in enumerate(rows_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                p = cell.paragraphs[0]
                set_paragraph_formatted_text(p, cell_text)
                for par in cell.paragraphs:
                    par.paragraph_format.space_after = Pt(3)
    doc.add_paragraph()


def main():
    text = MD_PATH.read_text(encoding="utf-8")
    doc = Document()
    doc.add_heading("MintCheck QA Checklist", 0)
    intro = (
        "Use this checklist to run through the product before TestFlight, launch, or after major changes. "
        "Test on a real device for OBD and deep links; simulator is fine for auth, settings, and navigation."
    )
    doc.add_paragraph(intro)
    doc.add_paragraph()

    current_heading = None
    table_lines = []
    notes_started = False
    notes_lines = []

    for line in text.splitlines():
        if line.strip() == "---":
            if table_lines:
                add_table_from_md_lines(doc, table_lines)
                table_lines = []
            continue
        if line.startswith("## "):
            if table_lines:
                add_table_from_md_lines(doc, table_lines)
                table_lines = []
            title = line[3:].strip()
            if title == "Notes":
                notes_started = True
                doc.add_heading("Notes", level=1)
                continue
            doc.add_heading(title, level=1)
            current_heading = title
            continue
        if notes_started:
            stripped = line.strip()
            if stripped.startswith("- **"):
                notes_lines.append(stripped.lstrip("- ").strip())
            continue
        if line.strip().startswith("|"):
            table_lines.append(line)
            continue
        if table_lines and line.strip() == "":
            add_table_from_md_lines(doc, table_lines)
            table_lines = []

    if table_lines:
        add_table_from_md_lines(doc, table_lines)

    if notes_lines:
        for nl in notes_lines:
            p = doc.add_paragraph(style="List Bullet")
            set_paragraph_formatted_text(p, nl)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    main()
